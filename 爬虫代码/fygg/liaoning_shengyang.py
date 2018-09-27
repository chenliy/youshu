#!/usr/bin/env python
# -*- encoding: utf-8 -*-
# Created on 2018/7/26 13:40
# Project: 
# @Author: ZQJ
# @Email : zihe@yscredit.com

from docx import Document
import pandas as pd
path = r'C:\Users\ll\Downloads\1531900336512579106.docx'
f = Document(path)
t = f.tables #这是一个列表，第0张表，第1张表....

print(len(t))
table = t[0]

len_row = len(table.rows)
print(len_row)

len_loc = len(table.columns)
print(len_loc)

title = []
for i in range(len_row):
    for j in range(len_loc):
        x = table.cell(i,j).text
        if x not in title:
            title.append(x)
    if len(title) == i+1:
        continue
    else:
        break
print(title)
print(i)

def get_title(path):

    f = Document(path)
    t = f.tables  # 这是一个列表，第0张表，第1张表....
    table = t[0]
    len_row = len(table.rows)
    len_loc = len(table.columns)
    title = []
    for i in range(len_row):
        for j in range(len_loc):
            x = table.cell(i,j).text
            if x not in title:
                title.append(x)
        if len(title) == i + 1:
            continue
        else:
            break
    title = title[:i]
    return

# for row in table.rows:
#     print(type(row))#docx.table._Row
#     print(type(row.cells))#tuple
#     print(len(row.cells))

    # for cell in row.cells:
    #     print(cell.text)
