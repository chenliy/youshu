#!/usr/bin/env python
# -*- encoding: utf-8 -*-
# Created on 2018/7/26 10:43
# Project: 
# @Author: ZQJ
# @Email : zihe@yscredit.com

##python docx学习
#安装#pip install python-docx

from docx import Document
from docx.shared import Inches
document = Document()#打开一个word,不加路径就是一个空的文档

#加入一个段，一句话

#这个方法返回对文档末尾新添加段落
paragraph = document.add_paragraph('你好啊')

#使用游标插入,允许插入在段落中间，在段落前面插入
prior_paragraph = paragraph.insert_paragraph_before('你好')

#插入标题,某人是一级标题，是在最前面的
document.add_heading('zhaoqijiong')

document.add_heading('chenlei',level=2)
#如果设置为0，则是标题段落


#添加分页，就算一页没满也强行翻页
document.add_page_break()

#添加表格,添加2行2列的表格
table = document.add_table(rows=2,cols=2)
#填充表格,用坐标的形式填充
cell = table.cell(0,1)
cell.text = 'zhao'

#也可以一列列，一行行的填
row = table.rows[1]
row.cells[0].text = 'qian'
row.cells[1].text = 'sun'

for row in table.rows:
    for cell in row.cells:
        print(cell.text)

document.save(r'C:\Users\ll\Downloads\demo.docx')